from pathlib import Path
from docx import Document

def delete_sections(document_path, output_path):
    
    doc = Document(document_path)
    
    titles_to_delete = ["References"]

    def remove_paragraph(paragraph):
        
        p = paragraph._element
        p.getparent().remove(p)
        p._element = p._p = None

    def remove_content_after_heading(doc, heading_text, delete_to_next_heading=True):
        
        delete = False
        for paragraph in doc.paragraphs:

            if heading_text in paragraph.text and paragraph.style.name.startswith("Heading"):
                delete = True
                remove_paragraph(paragraph)
                continue

            if delete:
                if delete_to_next_heading and paragraph.style.name.startswith("Heading"):
                    if heading_text not in paragraph.text and heading_text.lower() not in paragraph.text:
                        break
                remove_paragraph(paragraph)

    delete_contents = False
    for paragraph in doc.paragraphs:
        if "Contents" in paragraph.text:
            delete_contents = True
        if delete_contents:
            if paragraph.style.name.startswith("Heading") and "Contents" not in paragraph.text:
                break 
            remove_paragraph(paragraph)

    for title in titles_to_delete:
        remove_content_after_heading(doc, title)

    found_annex = False
    paragraphs_to_delete = []
    for i, paragraph in enumerate(doc.paragraphs):
        if "Annex" in paragraph.text and paragraph.style.name.startswith("Heading"):
            found_annex = True
        if found_annex:
            paragraphs_to_delete.append(paragraph)

    for paragraph in paragraphs_to_delete:
        remove_paragraph(paragraph)

    for table in doc.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)

    doc.save(output_path)

def batch_process(input_dir: str, output_dir: str):
    in_root = Path(input_dir)
    out_root = Path(output_dir)
    out_root.mkdir(parents=True, exist_ok=True)

    count = 0
    all_count = 0
    all_files = list(in_root.rglob("*.docx"))

    for in_file in all_files:
        # macOS metadata / Word temp 제외
        if in_file.name.startswith("._") or in_file.name.startswith("~$"):
            all_count += 1
            continue
        # __MACOSX 폴더
        if any(part == "__MACOSX" for part in in_file.parts):
            all_count += 1
            continue
        # 임시파일/숨김파일 스킵: ~$
        if in_file.name.startswith("~$"):
            all_count += 1
            continue

        # 입력 폴더 구조를 출력 폴더에 그대로 재현
        rel = in_file.relative_to(in_root)
        out_file = out_root / rel
        out_file.parent.mkdir(parents=True, exist_ok=True)

        try:
            delete_sections(str(in_file), str(out_file))
            all_count += 1
            count += 1
            print(f"[OK] {in_file} -> {out_file} | {all_count}/{len(all_files)}")
        except Exception as e:
            all_count += 1
            print(f"[FAIL] {in_file}: {e}")
            with open("bad_remove_content.txt", "a", encoding="utf-8") as f:
                f.write(str(in_file) + "\n")

    print(f"done. processed={count}/{len(all_files)}")

if __name__ == "__main__":
    batch_process(
        input_dir=r"\3gpp_docx", 
        output_dir=r"\3gpp_docx_cleaned"
    )